from langchain_community.document_loaders import PyPDFLoader
from docx2pdf import convert


def read_pdf(file_path, source=None):
    '''returns a list of pages from the pdf file
    each page has two props: page_content and metadata.
    metadata is a dictionary containing two keys: page: indicating the page number and source: indicating the source of the page
    '''
    try:
        loader = PyPDFLoader(file_path)
        pages = loader.load_and_split()

        # merge pages that are unmerged
        real_pages = []
        last_page_index = -1
        for page in pages:
            page_index = page.metadata['page']
            if page_index == last_page_index:
                real_pages[-1].page_content += page.page_content
            else:
                real_pages.append(page)
                last_page_index = page_index

        if source is None:
            # transform into dict
            real_pages = [page.to_dict() for page in real_pages]
            return real_pages

        else:
            final_results = []
            for page in real_pages:
                final_results.append({
                    'page_content': page.page_content,
                    'metadata': {
                        'page': str(page.metadata['page']),
                        'source': source
                    }
                })
            return final_results

    except Exception as e:
        print(f'Error reading pdf {file_path}: {e}')
        return []

import os
from docx import Document

def read_docx(file_path, source=None):
    '''returns a list of pages from the pdf file
    each page has two props: page_content and metadata.
    metadata is a dictionary containing two keys: page: indicating the page number and source: indicating the source of the page
    '''
    if source is None:
        source = file_path
    try:
        final_content = ''
        doc = Document(file_path)
        for para in doc.paragraphs:
            final_content += para.text + '\n'

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    final_content += cell.text + '\n'

        #split every 1000 words
        final_content_split = final_content.split()
        pages = []
        current_page = ''
        for word in final_content_split:
            if len(current_page.split()) < 1000:
                current_page += word + ' '
            else:
                pages.append(current_page)
                current_page = ''
        if current_page:
            pages.append(current_page)
        final_content = pages

        final_result = []
        for i, page in enumerate(final_content):
            final_result.append({
                'page_content': page,
                'metadata': {
                    'page': str(i),
                    'source': source
                }})

        return final_result
    except Exception as e:
        print(f'Error reading docx {source}: {e}')
        return []

# import pandas as pd
#
# def read_xlsx(file_path, source= None):
#     '''Returns a list of sheets from the Excel file.
#     Each sheet is represented as a dictionary with two properties: page_content and metadata.
#     Metadata is a dictionary containing two keys: page (indicating the sheet number) and source (indicating the source of the page).
#     '''
#     try:
#         # Load the Excel file
#         xls = pd.ExcelFile(file_path)
#         sheets = xls.sheet_names
#         final_results = []
#         if source is None:
#             source = file_path
#
#         # Iterate through each sheet in the Excel file
#         for i, sheet_name in enumerate(sheets):
#             # Read the sheet
#             sheet_df = pd.read_excel(xls, sheet_name)
#             # Convert the sheet data to a string
#             final_content = sheet_df.to_string()
#             # Append the sheet content and metadata to the results list
#             final_results.append({
#                 'page_content': final_content,
#                 'metadata': {
#                     'page': str(i + 1),  # Sheet number (1-indexed)
#                     'source': source,
#                     'sheet_name': sheet_name  # Include sheet name in metadata
#                 }
#             })
#
#         return final_results
#     except Exception as e:
#         print(f'Error reading xls {file_path}: {e}')
#         return []

def read_docx_pdf(file_path, source=None):
    if file_path.endswith('.docx'):
        file_path_pdf = file_path.replace('.docx', '.pdf')
        convert(file_path, file_path_pdf)
        return read_pdf(file_path_pdf, source)
    elif file_path.endswith('.pdf'):
        return read_pdf(file_path, source)
    else:
        print(f'Unsupported file type: {file_path}')
        return []



if __name__ == '__main__':


    print(read_pdf('testdata/Janvier2020_FAQCahierDesChargesFR.pdf'))
    print(len(read_pdf('testdata/Janvier2020_FAQCahierDesChargesFR.pdf')))

    print('-------------------')

    files = os.listdir('testdata')
    docx_files = [f for f in files if f.endswith('.docx')]
    # print(docx_files)


    xls_files = [f for f in files if f.endswith('.xls') or f.endswith('.xlsx')]
    # print(xls_files)

    for f in files:
        # print('file:', f)
        myread = read_docx_pdf_xlsx(os.path.join( 'testdata', f))
        print(myread)
        print(len(myread))
        print('-------------------')





